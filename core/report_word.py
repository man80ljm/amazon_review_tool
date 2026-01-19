# core/report_word.py
import os
from datetime import datetime
from typing import Optional, Dict, Any, List

import pandas as pd
from docx import Document
from docx.shared import Inches


# -----------------------------
# helpers
# -----------------------------
def _safe_makedirs(folder: str) -> None:
    if folder:
        os.makedirs(folder, exist_ok=True)


def _safe_exists(path: Optional[str]) -> bool:
    return bool(path) and os.path.exists(path)


def _pick_first_existing_col(df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
    for c in candidates:
        if c in df.columns:
            return c
    return None


def _add_picture_if_exists(doc: Document, path: Optional[str], heading: Optional[str] = None, width_in: float = 6.5):
    if _safe_exists(path):
        if heading:
            doc.add_heading(heading, level=2)
        try:
            doc.add_picture(path, width=Inches(width_in))
        except Exception:
            # 图片损坏/格式不支持也不应导致全流程失败
            doc.add_paragraph(f"[WARN] Failed to insert image: {path}")

def _add_df_table(
    doc: Document,
    df: pd.DataFrame,
    heading: str,
    max_rows: int = 30,
    max_cols: int = 10
):
    """把DataFrame以表格形式插入Word（防卡：限制行列）"""
    doc.add_heading(heading, level=2)

    if df is None or len(df) == 0:
        doc.add_paragraph("No data available.")
        return

    work = df.copy()

    # 限制列数（避免表太宽导致Word卡/乱）
    if work.shape[1] > max_cols:
        work = work.iloc[:, :max_cols].copy()
        doc.add_paragraph(f"[NOTE] Columns truncated to first {max_cols} columns.")

    # 限制行数（避免表太长）
    if len(work) > max_rows:
        work = work.head(max_rows).copy()
        doc.add_paragraph(f"[NOTE] Rows truncated to top {max_rows} rows.")

    # 建表：表头
    table = doc.add_table(rows=1, cols=work.shape[1])
    hdr = table.rows[0].cells
    for j, c in enumerate(work.columns):
        hdr[j].text = str(c)

    # 内容
    for _, r in work.iterrows():
        row = table.add_row().cells
        for j, c in enumerate(work.columns):
            v = r.get(c, "")
            try:
                if pd.isna(v):
                    v = ""
            except Exception:
                pass
            row[j].text = str(v)

def _topn_global(pivot: pd.DataFrame, n: int = 10) -> pd.DataFrame:
    """pivot: index=ASIN, columns=Attribute -> 返回全局TopN（按列均值降序）"""
    if pivot is None or len(pivot) == 0:
        return pd.DataFrame()
    s = pivot.mean(axis=0).sort_values(ascending=False).head(n)
    out = pd.DataFrame({
        "attribute": s.index.astype(str),
        "mean_value": s.values
    })
    return out

def _topn_per_asin(pivot: pd.DataFrame, n: int = 5) -> pd.DataFrame:
    """pivot: index=ASIN, columns=Attribute -> 返回每个ASIN TopN（长表：asin/attribute/value/rank）"""
    if pivot is None or len(pivot) == 0:
        return pd.DataFrame()
    rows = []
    for asin in pivot.index:
        row = pivot.loc[asin].sort_values(ascending=False).head(n)
        for rank, (attr, val) in enumerate(row.items(), start=1):
            rows.append({
                "asin": asin,
                "rank": rank,
                "attribute": str(attr),
                "value": float(val)
            })
    return pd.DataFrame(rows)

def _add_key_findings(
    doc: Document,
    g_share: Optional[pd.DataFrame] = None,   # columns: attribute, mean_value
    g_pain: Optional[pd.DataFrame] = None,    # columns: attribute, mean_value
    per_pain: Optional[pd.DataFrame] = None,  # columns: asin, rank, attribute, value
    opp: Optional[pd.DataFrame] = None,       # columns: asin, attribute, delta, pain(optional)
    topk: int = 3,
    show_metrics: bool = False,
    make_tables: bool = True,                # ✅ 新增：是否生成表格
):
    """
    把表格结果提炼成可引用的结论句（bullet points）。
    show_metrics=True 时，会把 mean_value/value/delta 一起带上。
    make_tables=True 时，会生成两张 Key Findings 表（全局/产品层）。
    """
    doc.add_heading("Key Findings / 关键发现", level=3)

    bullets = []

    def _fmt_num(x, nd=3):
        try:
            return f"{float(x):.{nd}f}"
        except Exception:
            return None

    def _fmt_share_percent(v):
        """
        自动判断 share 是比例(0~1)还是百分数(0~100)，避免出现 1639.6% 这种双乘问题。
        规则：>1.5 基本不可能是比例 -> 认为已经是百分数
        """
        try:
            x = float(v)
        except Exception:
            return None
        if x > 1.5:
            return f"{x:.1f}%"
        else:
            return f"{x*100.0:.1f}%"

    # -----------------------------
    # 1) Global pain TopK
    # -----------------------------
    if g_pain is not None and len(g_pain) > 0 and "attribute" in g_pain.columns:
        top = g_pain.head(topk).copy()
        if show_metrics and "mean_value" in top.columns:
            parts = []
            for _, r in top.iterrows():
                mv = _fmt_num(r.get("mean_value"), nd=3)
                if mv is None:
                    parts.append(str(r.get("attribute")))
                else:
                    parts.append(f"{r.get('attribute')} (mean_pain={mv})")
            bullets.append(f"Global pain Top{topk}: " + "; ".join(parts) + ".")
        else:
            top_attrs = top["attribute"].astype(str).tolist()
            bullets.append(f"Global pain Top{topk}: {', '.join(top_attrs)}.")

    # -----------------------------
    # 2) Global share TopK  ✅ 修复百分号
    # -----------------------------
    if g_share is not None and len(g_share) > 0 and "attribute" in g_share.columns:
        top = g_share.head(topk).copy()
        if show_metrics and "mean_value" in top.columns:
            parts = []
            for _, r in top.iterrows():
                pct = _fmt_share_percent(r.get("mean_value"))
                if pct is None:
                    parts.append(str(r.get("attribute")))
                else:
                    parts.append(f"{r.get('attribute')} ({pct})")
            bullets.append(f"Global share Top{topk}: " + "; ".join(parts) + ".")
        else:
            top_attrs = top["attribute"].astype(str).tolist()
            bullets.append(f"Global share Top{topk}: {', '.join(top_attrs)}.")

    # -----------------------------
    # 3) Opportunities TopK
    # -----------------------------
    if opp is not None and len(opp) > 0:
        asin_col = "asin" if "asin" in opp.columns else ("ASIN" if "ASIN" in opp.columns else None)
        attr_col = "attribute" if "attribute" in opp.columns else None
        delta_col = "delta" if "delta" in opp.columns else None

        if asin_col and attr_col:
            lines = []
            for _, r in opp.head(topk).iterrows():
                asin_v = str(r.get(asin_col, "-"))
                attr_v = str(r.get(attr_col, "-"))

                if show_metrics and delta_col:
                    dv = _fmt_num(r.get(delta_col), nd=3)
                    if dv is not None:
                        lines.append(f"{asin_v} → {attr_v} (delta={dv})")
                    else:
                        lines.append(f"{asin_v} → {attr_v}")
                else:
                    lines.append(f"{asin_v} → {attr_v}")

            bullets.append("Top opportunities (relative gaps): " + "; ".join(lines) + ".")

    # -----------------------------
    # 4) Per-ASIN primary pain (rank=1)
    # -----------------------------
    if per_pain is not None and len(per_pain) > 0 and {"asin", "rank", "attribute"}.issubset(set(per_pain.columns)):
        try:
            top1 = per_pain[per_pain["rank"] == 1].copy().head(5)
            lines = []
            for _, r in top1.iterrows():
                asin_v = str(r.get("asin"))
                attr_v = str(r.get("attribute"))
                if show_metrics and "value" in top1.columns:
                    vv = _fmt_num(r.get("value"), nd=3)
                    if vv is not None:
                        lines.append(f"{asin_v}: {attr_v} (pain={vv})")
                    else:
                        lines.append(f"{asin_v}: {attr_v}")
                else:
                    lines.append(f"{asin_v}: {attr_v}")
            if lines:
                bullets.append("Per-ASIN primary pain point: " + "; ".join(lines) + ".")
        except Exception:
            pass

    # 输出 bullet
    if not bullets:
        doc.add_paragraph("[NOTE] No sufficient data to summarize key findings.")
        return

    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # ----------------------------------------------------
    # ✅ 可选：生成 Key Findings 表格（更论文友好）
    # 依赖你已有的 _add_df_table(doc, df, heading=...)
    # ----------------------------------------------------
    if not make_tables:
        return

    try:
        # Table 1: Global
        rows1 = []
        if g_pain is not None and len(g_pain) > 0 and "attribute" in g_pain.columns:
            for i, r in g_pain.head(topk).iterrows():
                mv = r.get("mean_value", None)
                mv_s = _fmt_num(mv, nd=3) if (mv is not None) else None
                metric = f"mean_pain={mv_s}" if (show_metrics and mv_s is not None) else ""
                rows1.append(["Global Pain", int(i) + 1, str(r.get("attribute")), metric])

        if g_share is not None and len(g_share) > 0 and "attribute" in g_share.columns:
            for i, r in g_share.head(topk).iterrows():
                mv = r.get("mean_value", None)
                pct = _fmt_share_percent(mv) if (mv is not None) else None
                metric = f"mean_share={pct}" if (show_metrics and pct is not None) else ""
                rows1.append(["Global Share", int(i) + 1, str(r.get("attribute")), metric])

        df_global = pd.DataFrame(rows1, columns=["Type", "Rank", "Attribute", "Metric"])
        if len(df_global) > 0:
            _add_df_table(
                doc,
                df_global,
                heading="Key Findings Table (Global) / 关键发现表（全局）",
                max_rows=50,
                max_cols=6
            )

        # Table 2: Product-level
        rows2 = []
        if opp is not None and len(opp) > 0:
            asin_col = "asin" if "asin" in opp.columns else ("ASIN" if "ASIN" in opp.columns else None)
            attr_col = "attribute" if "attribute" in opp.columns else None
            delta_col = "delta" if "delta" in opp.columns else None
            if asin_col and attr_col:
                for _, r in opp.head(topk).iterrows():
                    asin_v = str(r.get(asin_col, "-"))
                    attr_v = str(r.get(attr_col, "-"))
                    dv = _fmt_num(r.get(delta_col), nd=3) if (show_metrics and delta_col) else None
                    metric = f"delta={dv}" if (dv is not None) else ""
                    rows2.append([asin_v, "Opportunity", attr_v, metric])

        if per_pain is not None and len(per_pain) > 0 and {"asin", "rank", "attribute"}.issubset(set(per_pain.columns)):
            top1 = per_pain[per_pain["rank"] == 1].copy().head(10)
            for _, r in top1.iterrows():
                asin_v = str(r.get("asin"))
                attr_v = str(r.get("attribute"))
                vv = _fmt_num(r.get("value"), nd=3) if (show_metrics and "value" in top1.columns) else None
                metric = f"pain={vv}" if (vv is not None) else ""
                rows2.append([asin_v, "Primary Pain", attr_v, metric])

        df_prod = pd.DataFrame(rows2, columns=["ASIN", "Finding Type", "Attribute", "Metric"])
        if len(df_prod) > 0:
            _add_df_table(
                doc,
                df_prod,
                heading="Key Findings Table (Product) / 关键发现表（产品层）",
                max_rows=80,
                max_cols=6
            )

    except Exception as e:
        doc.add_paragraph(f"[WARN] Key Findings table generation failed: {e}")

# -----------------------------
# sections
# -----------------------------
def _add_data_overview(doc: Document, df_all: pd.DataFrame, df_work: pd.DataFrame):
    doc.add_heading("Data Overview / 数据概览", level=2)
    doc.add_paragraph(f"Total rows (raw): {len(df_all)}")
    doc.add_paragraph(f"Rows used for clustering (after filter): {len(df_work)}")

    # show key columns if present
    cols_preview = [c for c in ["_text", "sentiment", "cluster_id", "ASIN", "Star"] if c in df_all.columns]
    if cols_preview:
        doc.add_paragraph(f"Key columns present: {', '.join(cols_preview)}")


def _add_sentiment_distribution(doc: Document, df_all: pd.DataFrame):
    if "sentiment" not in df_all.columns:
        return
    doc.add_heading("Sentiment Distribution / 情感分布", level=2)
    vc = df_all["sentiment"].value_counts(dropna=False).to_dict()
    doc.add_paragraph(str(vc))


def _add_k_table(doc: Document, k_to_inertia: Dict[int, float], k_to_sil: Dict[int, float], k_best: int):
    doc.add_heading("K Selection (Elbow + Silhouette) / K选择", level=2)
    doc.add_paragraph(f"Recommended K = {k_best}")

    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "K"
    hdr[1].text = "Inertia(SSE)"
    hdr[2].text = "Silhouette"

    ks = sorted(set(list(k_to_inertia.keys()) + list(k_to_sil.keys())))
    for k in ks:
        row = table.add_row().cells
        row[0].text = f"{k}{' (best)' if int(k) == int(k_best) else ''}"
        row[1].text = f"{float(k_to_inertia.get(k, 0.0)):.2f}"
        row[2].text = f"{float(k_to_sil.get(k, 0.0)):.4f}"


def _add_cluster_summary(doc: Document, summary_df: Optional[pd.DataFrame]):
    doc.add_heading("Cluster Summary / 聚类摘要", level=2)

    if summary_df is None or len(summary_df) == 0:
        doc.add_paragraph("No cluster summary available.")
        return

    # 兼容列名：只要能拿到 cluster_id/size/ratio/keywords 就行
    cid_col = _pick_first_existing_col(summary_df, ["cluster_id", "cluster", "cid"])
    size_col = _pick_first_existing_col(summary_df, ["cluster_size", "size", "count", "n"])
    ratio_col = _pick_first_existing_col(summary_df, ["ratio", "percent", "pct"])
    kw_col = _pick_first_existing_col(summary_df, ["keywords", "top_keywords", "kw"])

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "cluster_id"
    hdr[1].text = "cluster_size"
    hdr[2].text = "ratio"
    hdr[3].text = "keywords"

    for _, r in summary_df.iterrows():
        row = table.add_row().cells
        row[0].text = str(r.get(cid_col, ""))
        row[1].text = str(int(r.get(size_col, 0))) if size_col else ""
        try:
            row[2].text = f"{float(r.get(ratio_col, 0.0)):.3f}" if ratio_col else ""
        except Exception:
            row[2].text = str(r.get(ratio_col, "")) if ratio_col else ""
        row[3].text = str(r.get(kw_col, "")) if kw_col else ""


def _add_representatives(doc: Document, reps_df: Optional[pd.DataFrame], max_per_cluster: int = 3):
    doc.add_heading("Representative Reviews / 代表评论", level=2)

    if reps_df is None or len(reps_df) == 0:
        doc.add_paragraph("No representative reviews available.")
        return

    if "cluster_id" not in reps_df.columns:
        doc.add_paragraph("Missing column: cluster_id in reps_df.")
        return

    # 兼容 rank 字段
    rank_col = None
    for c in ["rank_in_cluster", "rank", "rank_idx", "order"]:
        if c in reps_df.columns:
            rank_col = c
            break

    # 兼容文本字段
    text_col = _pick_first_existing_col(reps_df, ["_text", "review_text", "text", "content", "评论内容"])

    # 兼容 asin/star/id 字段
    asin_col = _pick_first_existing_col(reps_df, ["ASIN", "asin", "_group"])
    star_col = _pick_first_existing_col(reps_df, ["Star", "star", "rating", "_score"])
    id_col = _pick_first_existing_col(reps_df, ["review_id", "id", "_id"])

    # cluster 顺序
    cluster_ids = sorted(pd.Series(reps_df["cluster_id"]).dropna().unique().tolist())

    for cid in cluster_ids:
        doc.add_heading(f"Cluster {cid}", level=3)
        sub = reps_df[reps_df["cluster_id"] == cid].copy()

        if rank_col:
            try:
                sub = sub.sort_values(rank_col)
            except Exception:
                pass

        sub = sub.head(max_per_cluster)

        for _, r in sub.iterrows():
            asin_v = r.get(asin_col, "-") if asin_col else "-"
            star_v = r.get(star_col, "-") if star_col else "-"
            id_v = r.get(id_col, "-") if id_col else "-"

            doc.add_paragraph(f"- (ASIN={asin_v}, Star={star_v}, id={id_v})")

            text = str(r.get(text_col, "")) if text_col else ""
            text = (text or "").strip()
            if not text:
                doc.add_paragraph("[EMPTY TEXT]")
            else:
                # 避免 Word 超长段落导致卡
                doc.add_paragraph(text[:900] + ("..." if len(text) > 900 else ""))

def _add_asin_attribute_section(
    doc: Document,
    asin_attr_xlsx: Optional[str],
    asin_attr_share_png: Optional[str],
    asin_attr_pain_png: Optional[str],
    key_findings_with_metrics: bool = False,   # ✅ 新增
):
    """
    核心升级：ASIN×Attribute（设计属性层）
    - 读 Excel: asin_attribute_matrix.xlsx
      * attribute_taxonomy
      * opportunity_top
    - 插入两张热力图：share/pain
    """
    # 三者都不存在就不输出
    if not (_safe_exists(asin_attr_xlsx) or _safe_exists(asin_attr_share_png) or _safe_exists(asin_attr_pain_png)):
        return

    doc.add_heading("ASIN × Attribute Insights (Design Attributes) / 跨ASIN设计属性洞察", level=2)

    # 1) taxonomy + opportunity（从xlsx读）
    if _safe_exists(asin_attr_xlsx):
        try:
            xls = pd.ExcelFile(asin_attr_xlsx)

            # 1.1 taxonomy
            if "attribute_taxonomy" in xls.sheet_names:
                tax = pd.read_excel(xls, sheet_name="attribute_taxonomy")
                _add_df_table(
                    doc,
                    tax,
                    heading="Attribute Taxonomy / 设计属性清单",
                    max_rows=40,
                    max_cols=8
                )
            else:
                doc.add_paragraph("[WARN] Sheet missing: attribute_taxonomy")

            # 1.2 opportunity
            if "opportunity_top" in xls.sheet_names:
                opp = pd.read_excel(xls, sheet_name="opportunity_top")
                _add_df_table(
                    doc,
                    opp,
                    heading="Opportunity Insights (Top) / 机会点总结（Top）",
                    max_rows=25,
                    max_cols=10
                )
            else:
                doc.add_paragraph("[WARN] Sheet missing: opportunity_top")

            # 1.3 share/pain TopN（新增）
            share_pivot = None
            pain_pivot = None

            if "asin_attribute_share" in xls.sheet_names:
                share_pivot = pd.read_excel(xls, sheet_name="asin_attribute_share", index_col=0)
            else:
                doc.add_paragraph("[WARN] Sheet missing: asin_attribute_share")

            if "asin_attribute_pain" in xls.sheet_names:
                pain_pivot = pd.read_excel(xls, sheet_name="asin_attribute_pain", index_col=0)
            else:
                doc.add_paragraph("[WARN] Sheet missing: asin_attribute_pain")

            # 全局TopN（Across ASIN）
            if share_pivot is not None and len(share_pivot) > 0:
                g_share = _topn_global(share_pivot, n=10)
                _add_df_table(doc, g_share, heading="Top Attributes by Share (Global) / 全局Top属性（占比）", max_rows=15, max_cols=6)

            if pain_pivot is not None and len(pain_pivot) > 0:
                g_pain = _topn_global(pain_pivot, n=10)
                _add_df_table(doc, g_pain, heading="Top Attributes by Pain (Global) / 全局Top属性（痛点优先级）", max_rows=15, max_cols=6)

            # 每个ASIN TopN（Per ASIN）
            if pain_pivot is not None and len(pain_pivot) > 0:
                per_pain = _topn_per_asin(pain_pivot, n=5)
                _add_df_table(doc, per_pain, heading="Top Pain Attributes per ASIN / 各ASIN Top痛点属性", max_rows=80, max_cols=6)

            if share_pivot is not None and len(share_pivot) > 0:
                per_share = _topn_per_asin(share_pivot, n=5)
                _add_df_table(doc, per_share, heading="Top Share Attributes per ASIN / 各ASIN Top占比属性", max_rows=80, max_cols=6)

        except Exception as e:
            doc.add_paragraph(f"[WARN] Failed to read asin_attribute_matrix.xlsx: {e}")

    # 2) 两张热力图（share/pain）
    _add_picture_if_exists(
        doc,
        asin_attr_share_png,
        heading="ASIN × Attribute Share Heatmap / 属性强度热力图（占比）",
        width_in=6.5
    )
    _add_picture_if_exists(
        doc,
        asin_attr_pain_png,
        heading="ASIN × Attribute Pain Heatmap / 痛点优先级热力图",
        width_in=6.5
    )

    # 3) Key Findings（新增：自动总结）
    try:
        _add_key_findings(
            doc, g_share=g_share, g_pain=g_pain, per_pain=per_pain, opp=opp, 
            topk=3, show_metrics=key_findings_with_metrics
        )
    except Exception as e:
        doc.add_paragraph(f"[WARN] Key Findings generation failed: {e}")

# -----------------------------
# main
# -----------------------------
def build_offline_report(
    *,
    cfg,
    output_dir: str,
    df_all: pd.DataFrame,
    df_work: pd.DataFrame,
    k_to_inertia: Dict[int, float],
    k_to_silhouette: Dict[int, float],
    k_best: int,
    cluster_summary: Optional[pd.DataFrame],
    reps_df: Optional[pd.DataFrame],
    k_plot_png: Optional[str] = None,
    asin_heatmap_png: Optional[str] = None,
    priority_png: Optional[str] = None,
     # ✅ 新增：ASIN×Attribute（核心升级）
    asin_attr_xlsx: Optional[str] = None,
    asin_attr_share_png: Optional[str] = None,
    asin_attr_pain_png: Optional[str] = None,
    key_findings_with_metrics: bool = False,
) -> str:
    """
    生成离线 Word：不依赖任何 LLM，尽最大可能“永不崩溃”
    - cfg 缺字段：用默认值兜底
    - 图缺失：跳过
    - DF 缺列：降级输出
    """
    _safe_makedirs(output_dir)

    # ---- cfg fallback ----
    title = getattr(cfg, "report_title", "Review Analysis Report")
    author = getattr(cfg, "report_author", "")
    subtitle = getattr(cfg, "report_subtitle", "")
    filename = getattr(cfg, "report_filename_offline", "review_analysis_report.docx")

    # 容错：用户给了 .doc 之类
    if not str(filename).lower().endswith(".docx"):
        filename = str(filename) + ".docx"

    out_path = os.path.join(output_dir, filename)

    doc = Document()

    # Title
    doc.add_heading(title, level=1)
    if subtitle:
        doc.add_paragraph(subtitle)
    if author:
        doc.add_paragraph(f"Author: {author}")
    doc.add_paragraph(f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Overview
    _add_data_overview(doc, df_all, df_work)
    _add_sentiment_distribution(doc, df_all)

    # K selection
    _add_k_table(doc, k_to_inertia or {}, k_to_silhouette or {}, int(k_best))

    if _safe_exists(k_plot_png):
        doc.add_paragraph("Legend:")
        doc.add_paragraph("- Solid line: WCSS/Inertia (Elbow)")
        doc.add_paragraph("- Dashed line: Silhouette score")
        doc.add_paragraph("- Vertical line: recommended K")
    _add_picture_if_exists(doc, k_plot_png, heading="K Selection Plot / K选择图", width_in=6.5)

    # Cluster summary + reps
    _add_cluster_summary(doc, cluster_summary)
    _add_representatives(doc, reps_df, max_per_cluster=int(getattr(cfg, "top_representatives", 3)) or 3)

    # Optional sections
    _add_picture_if_exists(doc, asin_heatmap_png, heading="Cross-ASIN Comparison / 跨ASIN对比", width_in=6.5)
    _add_picture_if_exists(doc, priority_png, heading="Priority Ranking / 优先级排序", width_in=6.5)
    
    # ✅ 新增：ASIN×Attribute 核心升级（表 + 新热力图）
    _add_asin_attribute_section(
    doc,
    asin_attr_xlsx,
    asin_attr_share_png,
    asin_attr_pain_png,
    key_findings_with_metrics=key_findings_with_metrics
    )

    # Save (never crash)
    try:
        doc.save(out_path)
    except Exception as e:
        # 最后兜底：换个文件名再保存
        fallback = os.path.join(output_dir, "review_analysis_report_fallback.docx")
        try:
            doc.save(fallback)
            out_path = fallback
        except Exception:
            raise RuntimeError(f"Failed to save Word report: {e}")

    return out_path
